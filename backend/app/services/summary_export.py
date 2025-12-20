from __future__ import annotations

import io
import re
from typing import Iterable, Literal, Optional


def _normalize_export_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "\u00a0": " ",
        "\u200b": "",
        "“": '"',
        "”": '"',
        "’": "'",
        "‘": "'",
        "—": "-",
        "–": "-",
        "…": "...",
    }
    return "".join(replacements.get(ch, ch) for ch in text)


def _strip_markdown_inline(text: str) -> str:
    if not text:
        return ""
    # Links: [label](url) -> label
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Inline code: `code` -> code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Bold/italic markers
    text = text.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    return text


def _iter_markdown_lines(markdown_text: str) -> Iterable[tuple[str, str]]:
    """Yield (kind, text) tuples for a small markdown subset.

    Kinds: blank | heading | bullet | text
    """
    for raw in (markdown_text or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            yield ("blank", "")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            yield ("heading", _strip_markdown_inline(heading_match.group(2).strip()))
            continue

        bullet_match = re.match(r"^\s*(?:[-*]|\d+\.)\s+(.*)$", line)
        if bullet_match:
            yield ("bullet", _strip_markdown_inline(bullet_match.group(1).strip()))
            continue

        yield ("text", _strip_markdown_inline(line.strip()))


def build_summary_docx(
    *,
    summary_md: str,
    title: Optional[str] = None,
    metadata_lines: Optional[list[str]] = None,
) -> bytes:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX export") from exc

    doc = Document()
    if title:
        doc.add_heading(_normalize_export_text(title), level=1)

    for line in metadata_lines or []:
        doc.add_paragraph(_normalize_export_text(line))

    if metadata_lines:
        doc.add_paragraph("")

    for kind, text in _iter_markdown_lines(summary_md):
        normalized = _normalize_export_text(text)
        if kind == "blank":
            doc.add_paragraph("")
        elif kind == "heading":
            doc.add_heading(normalized, level=2)
        elif kind == "bullet":
            doc.add_paragraph(normalized, style="List Bullet")
        else:
            doc.add_paragraph(normalized)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_summary_pdf(
    *,
    summary_md: str,
    title: Optional[str] = None,
    metadata_lines: Optional[list[str]] = None,
    paper: Literal["letter", "a4"] = "letter",
) -> bytes:
    try:
        import fitz  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("PyMuPDF is required for PDF export") from exc

    page_rect = fitz.paper_rect(paper)
    margin = 54  # 0.75"
    usable_width = page_rect.width - (margin * 2)
    usable_height = page_rect.height - (margin * 2)

    fontname = "helv"
    title_size = 18
    meta_size = 10
    body_size = 11
    line_height = body_size * 1.35

    def wrap_words(text: str, max_width: float, fontsize: float) -> list[str]:
        words = text.split()
        if not words:
            return [""]
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if fitz.get_text_length(candidate, fontname=fontname, fontsize=fontsize) <= max_width:
                current = candidate
                continue
            if current:
                lines.append(current)
                current = word
            else:
                lines.append(word)
                current = ""
        if current:
            lines.append(current)
        return lines

    doc = fitz.open()
    page = doc.new_page(width=page_rect.width, height=page_rect.height)
    y = margin

    def ensure_room(required_height: float) -> None:
        nonlocal page, y
        if y + required_height <= margin + usable_height:
            return
        page = doc.new_page(width=page_rect.width, height=page_rect.height)
        y = margin

    if title:
        ensure_room(title_size * 1.6)
        page.insert_text(
            (margin, y + title_size),
            _normalize_export_text(title),
            fontsize=title_size,
            fontname=fontname,
        )
        y += title_size * 1.7

    for line in metadata_lines or []:
        ensure_room(meta_size * 1.3)
        page.insert_text(
            (margin, y + meta_size),
            _normalize_export_text(line),
            fontsize=meta_size,
            fontname=fontname,
        )
        y += meta_size * 1.4

    if metadata_lines:
        y += meta_size * 0.6

    bullet_indent = 16
    for kind, text in _iter_markdown_lines(summary_md):
        normalized = _normalize_export_text(text)
        if kind == "blank":
            y += line_height * 0.6
            continue

        if kind == "heading":
            ensure_room(line_height * 1.2)
            for wrapped in wrap_words(normalized, usable_width, body_size):
                ensure_room(line_height)
                page.insert_text(
                    (margin, y + body_size),
                    wrapped.upper() if len(wrapped) < 80 else wrapped,
                    fontsize=body_size,
                    fontname=fontname,
                )
                y += line_height
            y += line_height * 0.2
            continue

        if kind == "bullet":
            wrapped_lines = wrap_words(normalized, usable_width - bullet_indent, body_size)
            for idx, wrapped in enumerate(wrapped_lines):
                ensure_room(line_height)
                if idx == 0:
                    page.insert_text(
                        (margin, y + body_size),
                        "-",
                        fontsize=body_size,
                        fontname=fontname,
                    )
                page.insert_text(
                    (margin + bullet_indent, y + body_size),
                    wrapped,
                    fontsize=body_size,
                    fontname=fontname,
                )
                y += line_height
            continue

        for wrapped in wrap_words(normalized, usable_width, body_size):
            ensure_room(line_height)
            page.insert_text(
                (margin, y + body_size),
                wrapped,
                fontsize=body_size,
                fontname=fontname,
            )
            y += line_height

    return doc.tobytes()

